from docx import Document
from docx.oxml import OxmlElement
import re
from typing import Dict, List, Tuple
import json
from pathlib import Path

class DocAnalyzer:
    def __init__(self):
        self.nikud_patterns = {
            'nikud': r'[\u05B0-\u05BC\u05C1-\u05C2\u05C4-\u05C5\u05C7]',
            'hebrew': r'[\u0590-\u05FF]'
        }
        self.MIN_WORDS = 20  # מינימום מילים לפסקה משמעותית
        
    def _has_complex_bold(self, run) -> bool:
        """בדיקה אם יש הדגשת Complex Script (w:bCs)"""
        rPr = run._r.get_or_add_rPr()
        return bool(rPr.xpath('.//w:bCs'))
    
    def _has_regular_bold(self, run) -> bool:
        """בדיקה אם יש הדגשה רגילה"""
        return bool(run.bold or False)
    
    def _has_nikud(self, text: str) -> bool:
        """בדיקה אם יש ניקוד בטקסט"""
        return bool(re.search(self.nikud_patterns['nikud'], text))
    
    def _count_hebrew_words(self, text: str) -> int:
        """ספירת מילים בעברית בטקסט"""
        # מנקה רווחים מיותרים
        text = ' '.join(text.split())
        # מחפש מילים שמכילות לפחות אות עברית אחת
        words = re.findall(f"\\b[\\w]*{self.nikud_patterns['hebrew']}[\\w]*\\b", text)
        return len(words)
    
    def analyze_paragraph(self, paragraph) -> Dict:
        """ניתוח פסקה בודדת"""
        result = {
            'text': '',
            'runs': [],
            'has_nikud': False,
            'nikud_count': 0,
            'word_count': 0,
            'is_significant': False,
            'bold_analysis': {
                'complex_bold': False,
                'regular_bold': False,
                'both_bold': False
            }
        }
        
        for run in paragraph.runs:
            run_text = run.text
            has_complex = self._has_complex_bold(run)
            has_regular = self._has_regular_bold(run)
            has_nikud = self._has_nikud(run_text)
            nikud_count = len(re.findall(self.nikud_patterns['nikud'], run_text))
            
            run_info = {
                'text': run_text,
                'complex_bold': has_complex,
                'regular_bold': has_regular,
                'has_nikud': has_nikud,
                'nikud_count': nikud_count
            }
            
            result['runs'].append(run_info)
            result['text'] += run_text
            result['has_nikud'] |= has_nikud
            result['nikud_count'] += nikud_count
            result['bold_analysis']['complex_bold'] |= has_complex
            result['bold_analysis']['regular_bold'] |= has_regular
            result['bold_analysis']['both_bold'] |= (has_complex and has_regular)
        
        # ספירת מילים וקביעה אם הפסקה משמעותית
        result['word_count'] = self._count_hebrew_words(result['text'])
        result['is_significant'] = result['word_count'] >= self.MIN_WORDS
            
        return result
    
    def analyze_document(self, doc_path: str, start_para: int = 0, end_para: int = None) -> Dict:
        """ניתוח מסמך שלם או חלק ממנו"""
        doc = Document(doc_path)
        paragraphs = doc.paragraphs[start_para:end_para]
        
        result = {
            'total_paragraphs': len(paragraphs),
            'significant_paragraphs': 0,
            'paragraphs_with_nikud': 0,
            'total_nikud': 0,
            'bold_stats': {
                'complex_bold': 0,
                'regular_bold': 0,
                'both_bold': 0
            },
            'paragraphs': []
        }
        
        for para in paragraphs:
            if not para.text.strip():  # דילוג על פסקאות ריקות
                continue
                
            para_analysis = self.analyze_paragraph(para)
            
            # מוסיף רק פסקאות משמעותיות לתוצאה
            if para_analysis['is_significant']:
                result['paragraphs'].append(para_analysis)
                result['significant_paragraphs'] += 1
                
                if para_analysis['has_nikud']:
                    result['paragraphs_with_nikud'] += 1
                result['total_nikud'] += para_analysis['nikud_count']
                
                if para_analysis['bold_analysis']['complex_bold']:
                    result['bold_stats']['complex_bold'] += 1
                if para_analysis['bold_analysis']['regular_bold']:
                    result['bold_stats']['regular_bold'] += 1
                if para_analysis['bold_analysis']['both_bold']:
                    result['bold_stats']['both_bold'] += 1
        
        return result

    def _get_sample_paragraphs(self, paragraphs: List[Dict], num_samples: int = 3) -> List[Dict]:
        """Get sample paragraphs from the analysis results"""
        if len(paragraphs) <= num_samples:
            return paragraphs
            
        # Get first, middle and last paragraph
        indices = [0, len(paragraphs)//2, -1]
        return [paragraphs[i] for i in indices]

    def print_analysis(self, analysis: Dict, show_full_text: bool = False):
        """הדפסת תוצאות הניתוח בפורמט קריא"""
        print(f"\nסיכום ניתוח המסמך:")
        print(f"==================")
        print(f"סה\"כ פסקאות משמעותיות (מעל {self.MIN_WORDS} מילים): {analysis['significant_paragraphs']}")
        print(f"פסקאות משמעותיות עם ניקוד: {analysis['paragraphs_with_nikud']}")
        print(f"סה\"כ סימני ניקוד בפסקאות משמעותיות: {analysis['total_nikud']}")
        print(f"\nסטטיסטיקת הדגשות בפסקאות משמעותיות:")
        print(f"- הדגשת Complex Script (w:bCs): {analysis['bold_stats']['complex_bold']}")
        print(f"- הדגשה רגילה: {analysis['bold_stats']['regular_bold']}")
        print(f"- שתי ההדגשות יחד: {analysis['bold_stats']['both_bold']}")
        
        if show_full_text:
            print("\nדוגמאות מייצגות מהפסקאות המשמעותיות:")
            sample_paragraphs = self._get_sample_paragraphs(analysis['paragraphs'])
            for i, para in enumerate(sample_paragraphs, 1):
                print(f"\nפסקה {i}:")
                print(f"טקסט ({para['word_count']} מילים): {para['text']}")
                print(f"ניקוד: {'יש' if para['has_nikud'] else 'אין'} ({para['nikud_count']} סימנים)")
                print("הדגשות:", end=" ")
                if para['bold_analysis']['complex_bold']:
                    print("Complex Script", end=" ")
                if para['bold_analysis']['regular_bold']:
                    print("Regular", end=" ")
                if para['bold_analysis']['both_bold']:
                    print("(שניהם יחד)", end=" ")
                print()

    def find_matching_paragraphs(self, source_doc: str, target_doc: str) -> List[Dict]:
        """מוצא פסקאות מקבילות בין מסמך המקור למסמך היעד"""
        source = Document(source_doc)
        target = Document(target_doc)
        
        matches = []
        
        # מנקה ומנרמל את הטקסטים ובודק מספר מילים
        source_texts = []
        for i, p in enumerate(source.paragraphs):
            text = p.text.strip()
            if text:
                word_count = self._count_hebrew_words(text)
                if word_count >= self.MIN_WORDS:
                    source_texts.append((i, self._normalize_text(text)))
        
        target_texts = []
        for i, p in enumerate(target.paragraphs):
            text = p.text.strip()
            if text:
                word_count = self._count_hebrew_words(text)
                if word_count >= self.MIN_WORDS:
                    target_texts.append((i, self._normalize_text(text)))
        
        # מוצא התאמות
        for s_idx, s_text in source_texts:
            for t_idx, t_text in target_texts:
                # בודק אם הטקסטים זהים לאחר נרמול
                if s_text == t_text:
                    matches.append({
                        'source': {
                            'index': s_idx,
                            'text': source.paragraphs[s_idx].text,
                            'normalized': s_text
                        },
                        'target': {
                            'index': t_idx,
                            'text': target.paragraphs[t_idx].text,
                            'normalized': t_text
                        }
                    })
                    break  # עובר לפסקה הבאה במקור
        
        return matches

def main():
    analyzer = DocAnalyzer()
    
    # ניתוח קובץ המקור
    print("\nמנתח את קובץ המקור...")
    source_analysis = analyzer.analyze_document('source.docx', start_para=0, end_para=100)  # בודק יותר פסקאות
    analyzer.print_analysis(source_analysis, show_full_text=True)
    
    # ניתוח קובץ היעד
    print("\nמנתח את קובץ היעד...")
    target_analysis = analyzer.analyze_document('target.docx', start_para=0, end_para=100)  # בודק יותר פסקאות
    analyzer.print_analysis(target_analysis, show_full_text=True)
    
    # שמירת התוצאות לקובץ
    with open('analysis_results.json', 'w', encoding='utf-8') as f:
        json.dump({
            'source': source_analysis,
            'target': target_analysis
        }, f, ensure_ascii=False, indent=2)

if __name__ == '__main__':
    main() 