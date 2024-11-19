import streamlit as st
import anthropic
import json
from prompt_template import SYSTEM_PROMPT, PROMPT_TEMPLATE
from examples import INTERPRETATION_EXAMPLES
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from services.docx_generator import create_interpretation_docx
from services.text_generator import create_interpretation_txt
import io

def get_interpretation(text):
    api_key = st.secrets["ANTHROPIC_API_KEY"]
    client = anthropic.Anthropic(api_key=api_key)
    
    message = client.messages.create(
        model="claude-3-5-haiku-latest",
        max_tokens=4096,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user", 
            "content": PROMPT_TEMPLATE.format(
                text_to_analyze=text
            )
        }]
    )
    
    # Extract JSON from response
    try:
        response_text = message.content[0].text
        json_start = response_text.find('```json\n') + 8
        json_end = response_text.find('```', json_start)
        json_str = response_text[json_start:json_end].strip()
        return json.loads(json_str)
    except Exception as e:
        st.error(f"Failed to parse response: {e}")
        return None

def set_rtl(paragraph):
    """Set paragraph direction to RTL"""
    pPr = paragraph._p.get_or_add_pPr()
    biDi = OxmlElement('w:bidi')
    pPr.insert_element_before(biDi, 'w:jc')

def create_interpretation_docx(interpretation):
    doc = Document()
    
    # Set RTL for document
    section = doc.sections[0]
    section._sectPr.xpath('./w:bidi')[0].set(qn('w:val'), '1')
    
    # Set font
    style = doc.styles['Normal']
    style.font.name = 'David'
    style.font.size = Pt(12)
    
    # Letter - centered
    letter_para = doc.add_paragraph()
    set_rtl(letter_para)
    letter_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    letter_para.add_run(interpretation["letter"])
    
    # Original text - right aligned, bold
    text_para = doc.add_paragraph()
    set_rtl(text_para)
    text_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    text_para.add_run(interpretation["original_text"]).bold = True
    
    # Two line breaks
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Difficult words in one paragraph
    words_para = doc.add_paragraph()
    set_rtl(words_para)
    words_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    for i, word in enumerate(interpretation["difficult_words"]):
        words_para.add_run(word["word"]).bold = True
        words_para.add_run(f" - {word['explanation']}")
        if i < len(interpretation["difficult_words"]) - 1:
            words_para.add_run(" ; ")
    
    # Two line breaks
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Detailed interpretation
    interp_para = doc.add_paragraph()
    set_rtl(interp_para)
    interp_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    for i, detail in enumerate(interpretation["detailed_interpretation"]):
        interp_para.add_run(detail["quote"]).bold = True
        interp_para.add_run(f" - {detail['explanation']}")
        if i < len(interpretation["detailed_interpretation"]) - 1:
            interp_para.add_run(" ; ")
        else:
            interp_para.add_run(".")
    
    return doc

def main():
    # Set basic page config
    st.set_page_config(
        layout="wide",
        initial_sidebar_state="collapsed",
        page_title="מנדי - עוזר אישי לכתיבת פירוש תורני",
        page_icon="📚"
    )
    
    # Minimal CSS just for RTL support
    st.markdown("""
        <style>
            .stApp { direction: rtl; }
            .stTextArea textarea { direction: rtl; }
        </style>
    """, unsafe_allow_html=True)

    st.title("מנדי - עוזר אישי לכתיבת פירוש תורני")
    st.write("הכנס קטע טקסט לניתוח ופירוש מעמיק")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        user_text = st.text_area("הכנס טקסט לפירוש:", height=200)
        if st.button("פרש את הטקסט"):
            if not user_text:
                st.error("אנא הכנס טקסט לפירוש")
                return
                
            with st.spinner("מנתח את הטקסט..."):
                interpretation = get_interpretation(user_text)
                
                if interpretation:
                    with col2:
                        st.subheader("טקסט מקורי")
                        st.write(interpretation["original_text"])
                        
                        st.subheader("אות")
                        st.write(interpretation["letter"])
                        
                        st.subheader("מילים קשות")
                        for word in interpretation["difficult_words"]:
                            st.write(f"**{word['word']}**: {word['explanation']}")
                        
                        st.subheader("פירוש מפורט")
                        for detail in interpretation["detailed_interpretation"]:
                            st.write(f"**ציטוט**: {detail['quote']}")
                            st.write(f"**פירוש**: {detail['explanation']}")
                            st.markdown("---")
                        
                        # Add download button for txt
                        text_content = create_interpretation_txt(interpretation)
                        
                        st.download_button(
                            label="הורד כקובץ טקסט",
                            data=text_content.encode('utf-8'),
                            file_name="interpretation.txt",
                            mime="text/plain"
                        )

if __name__ == "__main__":
    main() 