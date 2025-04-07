import os
import tempfile
import shutil
from pathlib import Path
from docx import Document
import logging
import re
import json
import datetime

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Import local modules
from services.document_processor import DocumentProcessor, Section
from services.nikud_service import NikudService
from services.gemini_service import GeminiService
from services.usage_logger import UsageLogger, streamlit_logger as st_log

class SectionMatcherTester:
    def __init__(self):
        # Initialize the required services
        self.doc_processor = DocumentProcessor()
        self.nikud_service = NikudService()  # This will use the updated GeminiService

        # Setup the output directory with timestamp
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.temp_dir = tempfile.mkdtemp(prefix=f"section_matcher_test_{timestamp}_")
        logger.info(f"Created temporary directory: {self.temp_dir}")

    def run_test(self, source_path="source.docx", target_path="target.docx"):
        """Run the section matching test on the source and target files"""
        logger.info(f"Processing files: {source_path} and {target_path}")

        # Step 1: Read the DOCX files
        logger.info("Reading DOCX files...")
        source_text, source_doc = self.nikud_service._read_docx(source_path)
        target_text, target_doc = self.nikud_service._read_docx(target_path)

        # Step 2: Split into sections
        logger.info("Splitting documents into sections...")
        source_sections = self.doc_processor.split_to_sections(source_text)
        target_sections = self.doc_processor.split_to_sections(target_text)

        logger.info(f"Found {len(source_sections)} sections in source document")
        logger.info(f"Found {len(target_sections)} sections in target document")

        # Save all sections overview
        self._write_sections_overview(source_sections, os.path.join(self.temp_dir, "source_sections.txt"))
        self._write_sections_overview(target_sections, os.path.join(self.temp_dir, "target_sections.txt"))

        # Step 3: Find matching sections
        logger.info("Finding matching sections...")
        matches = self.doc_processor.find_matching_sections(source_sections, target_sections)
        
        logger.info(f"Found {len(matches)} matching section pairs")

        # Create a map of all matches
        self._create_match_mapping(matches)

        # Step 4: Create subdirectories and save matched sections
        for i, (source_section, target_section) in enumerate(matches):
            # Create a subdirectory for this match
            match_dir = os.path.join(self.temp_dir, f"match_{i+1}")
            os.makedirs(match_dir, exist_ok=True)
            
            # Create source and target docx files
            self._create_section_docx(source_section, os.path.join(match_dir, "source.docx"))
            self._create_section_docx(target_section, os.path.join(match_dir, "target.docx"))
            
            # Create text files with section content
            self._write_text_file(source_section.content, os.path.join(match_dir, "source_content.txt"))
            self._write_text_file(target_section.content, os.path.join(match_dir, "target_content.txt"))
            
            # Create info file with match details
            self._write_match_info(i+1, source_section, target_section, match_dir)
            
            # Create the prepared nikud content
            prepared_content = self.doc_processor.prepare_for_nikud(source_section, target_section)
            if prepared_content:
                self._write_text_file(
                    json.dumps(prepared_content, ensure_ascii=False, indent=2),
                    os.path.join(match_dir, "prepared_content.json")
                )
            
            logger.info(f"Created match files in {match_dir}")

        # Create a summary file
        self._create_summary_report(matches)

        logger.info(f"All matches processed. Results saved in {self.temp_dir}")
        return self.temp_dir

    def _write_sections_overview(self, sections, output_path):
        """Write an overview of all sections to a file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Total sections: {len(sections)}\n\n")
            for i, section in enumerate(sections):
                f.write(f"Section {i+1}: {section.header}\n")
                f.write(f"{'='*40}\n")
                f.write(f"Content length: {len(section.content)} characters\n")
                f.write(f"First sentence: {section.first_sentence}\n")
                f.write(f"Main content length: {len(section.main_content) if section.main_content else 0} characters\n")
                f.write(f"Preview: {section.content[:100]}...\n\n")

    def _create_match_mapping(self, matches):
        """Create a JSON file mapping all matches"""
        mapping = []
        for i, (source_section, target_section) in enumerate(matches):
            mapping.append({
                "match_number": i + 1,
                "source_header": source_section.header,
                "target_header": target_section.header,
                "match_directory": f"match_{i+1}",
                "bold_count": self._count_bold_tags(target_section.content)
            })
        
        # Write the mapping to a JSON file
        mapping_path = os.path.join(self.temp_dir, "match_mapping.json")
        with open(mapping_path, 'w', encoding='utf-8') as f:
            json.dump(mapping, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Created match mapping file: {mapping_path}")

    def _create_summary_report(self, matches):
        """Create a human-readable summary report of all matches"""
        report_path = os.path.join(self.temp_dir, "summary_report.txt")
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(f"Summary Report - Section Matcher Test\n")
            f.write(f"{'='*50}\n\n")
            f.write(f"Test conducted on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"Total matches found: {len(matches)}\n\n")
            
            f.write("Match Summary:\n")
            f.write(f"{'-'*40}\n")
            
            total_bold_tags = 0
            for i, (source_section, target_section) in enumerate(matches):
                bold_count = self._count_bold_tags(target_section.content)
                total_bold_tags += bold_count
                f.write(f"Match {i+1}: {source_section.header} ↔️ {target_section.header}\n")
                f.write(f"  Bold tags: {bold_count}\n")
                f.write(f"  Source length: {len(source_section.content)} chars\n")
                f.write(f"  Target length: {len(target_section.content)} chars\n")
                f.write(f"{'-'*40}\n")
            
            f.write(f"\nTotal bold tags across all sections: {total_bold_tags}\n")
            f.write(f"\nResults directory: {self.temp_dir}\n")
        
        logger.info(f"Created summary report: {report_path}")

    def _create_section_docx(self, section: Section, output_path: str):
        """Create a DOCX file containing just the section content"""
        doc = Document()
        # Add section header as title
        doc.add_heading(section.header, level=1)
        # Add section content
        for paragraph in section.content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                # Add text with HTML tags intact
                p.add_run(paragraph)
        # Save the document
        doc.save(output_path)

    def _write_text_file(self, content: str, output_path: str):
        """Write content to a text file with UTF-8 encoding"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def _count_bold_tags(self, text: str) -> int:
        """Count number of bold tags in text"""
        return len(re.findall(r'<b>.*?</b>', text))

    def _write_match_info(self, match_num: int, source_section: Section, target_section: Section, match_dir: str):
        """Write match information to a text file"""
        info_path = os.path.join(match_dir, "match_info.txt")
        with open(info_path, 'w', encoding='utf-8') as f:
            f.write(f"Match {match_num}\n")
            f.write(f"{'='*40}\n\n")
            
            f.write(f"Source Section Header: {source_section.header}\n")
            f.write(f"Target Section Header: {target_section.header}\n\n")
            
            f.write(f"Source First Sentence: {source_section.first_sentence}\n")
            f.write(f"Target First Sentence: {target_section.first_sentence}\n\n")
            
            # Count bold tags in target
            bold_count = self._count_bold_tags(target_section.content)
            f.write(f"Bold words/sections in target: {bold_count}\n\n")
            
            # Source content stats
            f.write(f"Source content length: {len(source_section.content)} characters\n")
            f.write(f"Target content length: {len(target_section.content)} characters\n")

def main():
    """Main function to run the test"""
    tester = SectionMatcherTester()
    output_dir = tester.run_test()
    print(f"\nTest completed! Results saved to: {output_dir}")
    print("You can examine the matched sections in the subdirectories.")

if __name__ == "__main__":
    main() 