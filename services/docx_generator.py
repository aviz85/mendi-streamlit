from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Set paragraph direction to RTL"""
    pPr = paragraph._p.get_or_add_pPr()
    if not pPr.xpath('.//w:bidi'):  # Only add if not exists
        biDi = OxmlElement('w:bidi')
        pPr.insert_element_before(biDi, 'w:jc')

def create_interpretation_docx(interpretation):
    doc = Document()
    
    # Set RTL for document
    section = doc.sections[0]
    sectPr = section._sectPr
    if not sectPr.xpath('./w:bidi'):  # Only add if not exists
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        sectPr.append(bidi)
    
    # Set font
    style = doc.styles['Normal']
    style.font.name = 'David'
    style.font.size = Pt(12)
    
    # Letter
    letter_para = doc.add_paragraph()
    set_rtl(letter_para)
    letter_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    letter_para.add_run(interpretation["letter"])
    
    # Original text
    text_para = doc.add_paragraph()
    set_rtl(text_para)
    text_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    text_para.add_run(interpretation["original_text"])
    
    # Single line spacing
    doc.add_paragraph()
    
    # Difficult words
    if "difficult_words" in interpretation:
        words_para = doc.add_paragraph()
        set_rtl(words_para)
        words_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        for i, word in enumerate(interpretation["difficult_words"]):
            words_para.add_run(word["word"]).bold = True
            words_para.add_run(f" - {word['explanation']}")
            if i < len(interpretation["difficult_words"]) - 1:
                words_para.add_run(" ; ")
    
    # Single line spacing
    doc.add_paragraph()
    
    # Detailed interpretation
    if "detailed_interpretation" in interpretation:
        interp_para = doc.add_paragraph()
        set_rtl(interp_para)
        interp_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        for i, detail in enumerate(interpretation["detailed_interpretation"]):
            if "quote" in detail:  # Check if quote exists
                interp_para.add_run(detail["quote"]).bold = True
                interp_para.add_run(f" - {detail['explanation']}")
                if i < len(interpretation["detailed_interpretation"]) - 1:
                    interp_para.add_run(" ; ")
                else:
                    interp_para.add_run(".")
    
    return doc