import re
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import RGBColor
from rapidfuzz import fuzz
import logging
from difflib import SequenceMatcher, Match

logging.basicConfig(level=logging.INFO, format='%(message)s')
logger = logging.getLogger(__name__)

class NikudMapper:
    def __init__(self, bold_only: bool = False):
        self.bold_only = bold_only
        self.nikud_patterns = {
            'nikud': r'[\u05B0-\u05BC\u05C1-\u05C2\u05C4-\u05C5\u05C7]',
            'hebrew': r'[\u0590-\u05FF]'
        }
        self._nikud_cache = {}  # מטמון למילים מנוקדות
        self._clean_cache = {}  # מטמון למילים נקיות
        
    def _create_virtual_copy(self, text: str) -> Tuple[List[str], List[str]]:
        """
        יצירת עותק וירטואלי של המילים המנוקדות
        
        Returns:
            Tuple[List[str], List[str]]: (מילים מנוקדות, מילים ללא ניקוד)
        """
        # הסרת שורות חדשות ורווחים מיותרים
        text = ' '.join(text.split())
        
        # מציאת כל המילים בעברית עם הפיסוק שלהן
        matches = list(re.finditer(f"({self.nikud_patterns['hebrew']}+)([^\u0590-\u05FF]*)", text))
        nikud_words = []  # מילים מנוקדות
        clean_words = []  # מילים ללא ניקוד
        
        for match in matches:
            word = match.group(1)
            nikud_words.append(word)
            clean_words.append(self._strip_nikud(word))
        
        return nikud_words, clean_words

    def _find_overlap(self, source_words: List[str], target_text: str) -> Tuple[int, int, int, int]:
        """מציאת חפיפה בין מילות המקור לטקסט היעד"""
        # מציאת כל המילים בטקסט היעד
        target_matches = list(re.finditer(f"({self.nikud_patterns['hebrew']}+)", target_text))
        target_words = [m.group(1) for m in target_matches]
        
        # הסרת ניקוד מהמילים
        source_clean = [self._strip_nikud(word) for word in source_words]
        target_clean = [self._strip_nikud(word) for word in target_words]
        
        # מציאת החפיפה הטובה ביותר
        best_overlap = (0, 0, 0, 0)
        best_score = 0
        min_size = 3  # גודל מינימלי לחפיפה
        window_step = 1  # קפיצות של מילה אחת
        similarity_threshold = 85  # סף דמיון
        
        # חיפוש חפיפות בחלונות
        for size in range(min_size, min(len(target_clean) + 1, len(source_clean) + 1)):
            for j in range(0, len(target_clean) - size + 1, window_step):
                target_slice = ' '.join(target_clean[j:j+size])
                
                # חיפוש בטקסט המקור
                for i in range(0, len(source_clean) - size + 1, window_step):
                    source_slice = ' '.join(source_clean[i:i+size])
                    
                    # חישוב דמיון
                    ratio = fuzz.ratio(source_slice, target_slice)
                    if ratio < similarity_threshold:
                        continue
                    
                    # חישוב ציון
                    position_score = 1 - (abs(i - j) / max(len(source_clean), len(target_clean)))  # העדפה למיקום דומה
                    size_score = size / len(target_clean)  # העדפה לחפיפות ארוכות
                    ratio_score = ratio / 100  # דמיון טקסטואלי
                    
                    score = (
                        size_score * 0.5 +     # 50% משקל לאורך
                        ratio_score * 0.3 +    # 30% משקל לדמיון
                        position_score * 0.2    # 20% משקל למיקום
                    )
                    
                    if score > best_score:
                        best_score = score
                        best_overlap = (i, i+size, j, j+size)
                        
                        # אם מצאנו חפיפה מספיק טובה, נחזיר אותה מיד
                        if score > 0.9:
                            return best_overlap
        
        # אם לא מצאנו חפיפה טובה, נחפש חפיפה מילה-מילה
        if best_overlap == (0, 0, 0, 0):
            for i, source_word in enumerate(source_clean):
                for j, target_word in enumerate(target_clean):
                    if source_word == target_word:
                        return (i, i+1, j, j+1)
        
        return best_overlap

    def _calc_context_score(self, source_words: List[str], target_words: List[str], match: Match) -> float:
        """חישוב ציון הקשר לחפיפה"""
        # בדיקת מילים לפני ואחרי החפיפה
        context_size = 2  # כמה מילים לבדוק בכל צד
        
        before_source = source_words[max(0, match.a-context_size):match.a]
        after_source = source_words[match.a+match.size:match.a+match.size+context_size]
        
        before_target = target_words[max(0, match.b-context_size):match.b]
        after_target = target_words[match.b+match.size:match.b+match.size+context_size]
        
        # חישוב דמיון בין ההקשרים
        before_ratio = fuzz.ratio(' '.join(before_source), ' '.join(before_target)) / 100
        after_ratio = fuzz.ratio(' '.join(after_source), ' '.join(after_target)) / 100
        
        return (before_ratio + after_ratio) / 2

    def _align_cursors(self, source_words: List[str], target_words: List[str], 
                      start_source: int, start_target: int) -> Tuple[int, int]:
        """יישור סמנים לתחילת החפיפה"""
        source_cursor = start_source
        target_cursor = start_target
        
        source_word = source_words[source_cursor]
        target_word = target_words[target_cursor]
        
        logger.info(
            f"מיקום סמן מקור: {source_cursor} מורה על המילה - {source_word}\n"
            f"מיקום סמן יעד: {target_cursor} מורה על המילה - {target_word}"
        )
        
        return source_cursor, target_cursor

    def add_nikud_to_text(self, source_text: str, target_text: str) -> str:
        """הוספת ניקוד לטקסט"""
        # בדיקה במטמון
        cache_key = f"{hash(source_text)}:{hash(target_text)}"
        if cache_key in self._nikud_cache:
            return self._nikud_cache[cache_key]
            
        # יצירת עותק וירטואלי
        source_nikud, source_clean = self._create_virtual_copy(source_text)
        
        # מציאת חפיפה
        start_s, end_s, start_t, end_t = self._find_overlap(source_clean, target_text)
        
        # אם לא נמצאה חפיפה, מחזירים את הטקסט המקורי
        if start_s == end_s == start_t == end_t == 0:
            self._nikud_cache[cache_key] = target_text
            return target_text
        
        # מציאת כל המילים בטקסט היעד עם הפיסוק שלהן
        target_matches = list(re.finditer(f"({self.nikud_patterns['hebrew']}+)([^\u0590-\u05FF]*)", target_text))
        target_words = [m.group(1) for m in target_matches]  # המילים עצמן
        target_puncts = [m.group(2) for m in target_matches]  # הפיסוק אחרי כל מילה
        
        # עיבוד הטקסט לפי החפיפה שנמצאה
        result = []
        
        # מילים לפני החפיפה
        for i in range(start_t):
            result.append(target_words[i] + target_puncts[i])
        
        # מילים בתוך החפיפה
        for i in range(start_t, end_t):
            source_idx = start_s + (i - start_t)
            if source_idx < len(source_clean):
                target_clean = self._strip_nikud(target_words[i])
                source_clean_word = source_clean[source_idx]
                
                # בדיקה אם המילה הנוכחית מתאימה למילה במקור
                if target_clean == source_clean_word:
                    result.append(source_nikud[source_idx] + target_puncts[i])
                else:
                    # חיפוש מילה מנוקדת מתאימה בכל המקור
                    found = False
                    for j, (nikud_word, clean_word) in enumerate(zip(source_nikud, source_clean)):
                        if clean_word == target_clean:
                            result.append(nikud_word + target_puncts[i])
                            found = True
                            break
                    if not found:
                        result.append(target_words[i] + target_puncts[i])
            else:
                result.append(target_words[i] + target_puncts[i])
        
        # מילים אחרי החפיפה
        for i in range(end_t, len(target_words)):
            result.append(target_words[i] + target_puncts[i])
        
        # חיבור התוצאה למחרוזת אחת
        final_result = ' '.join(result)
        
        # שמירה במטמון
        self._nikud_cache[cache_key] = final_result
        return final_result

    def _has_nikud(self, text: str) -> bool:
        return bool(re.search(self.nikud_patterns['nikud'], text))

    def _has_hebrew(self, text: str) -> bool:
        return bool(re.search(self.nikud_patterns['hebrew'], text))

    def _strip_nikud(self, text: str) -> str:
        if text in self._clean_cache:
            return self._clean_cache[text]
        result = re.sub(self.nikud_patterns['nikud'], '', text)
        self._clean_cache[text] = result
        return result

    def _is_bold(self, word: str) -> bool:
        # TODO: לממש בדיקת הדגשה לפי המסמך
        return True

    def process_docx(self, source_path: str, target_path: str, output_path: str) -> None:
        """
        עיבוד קובץ Word והוספת ניקוד
        
        Args:
            source_path: נתיב לקובץ המקור המנוקד
            target_path: נתיב לקובץ היעד
            output_path: נתיב לקובץ הפלט
        """
        logger.info(f"מתחיל עיבוד קבצים:\nמקור: {source_path}\nקלט: {target_path}\nפלט: {output_path}")
        
        try:
            # קריאת קובץ המקור והכנת המילים המנוקדות
            logger.info("קורא קובץ מקור...")
            source_doc = Document(source_path)
            source_text = '\n'.join(p.text for p in source_doc.paragraphs)
            logger.info(f"טקסט מקור ({len(source_text)} תווים):\n{source_text[:200]}...")
            source_nikud, source_clean = self._create_virtual_copy(source_text)
            logger.info(f"נמצאו {len(source_nikud)} מילים מנוקדות במקור")
            
            # קריאת קובץ היעד
            logger.info("קורא קובץ יעד...")
            target_doc = Document(target_path)
            
            # עיבוד כל פסקה
            logger.info(f"מעבד {len(target_doc.paragraphs)} פסקאות...")
            for i, paragraph in enumerate(target_doc.paragraphs, 1):
                # מציאת חפיפה לכל הפסקה
                paragraph_text = paragraph.text
                if not self._has_hebrew(paragraph_text):
                    logger.info(f"פסקה {i}: אין טקסט בעברית, דילוג")
                    continue
                
                logger.info(f"פסקה {i} ({len(paragraph_text)} תווים):\n{paragraph_text}")
                
                # עיבוד הריצות בפסקה
                for j, run in enumerate(paragraph.runs, 1):
                    # בדיקת הדגשה לפי XML
                    rPr = run._r.get_or_add_rPr()
                    is_bold = bool(rPr.xpath('.//w:b') or rPr.xpath('.//w:bCs'))
                    
                    if is_bold or not self.bold_only:
                        original_text = run.text
                        if self._has_hebrew(original_text):
                            logger.info(f"מעבד ריצה {j} (מודגש={is_bold}): {original_text}")
                            run.text = self.add_nikud_to_text(source_text, original_text)
                            logger.info(f"אחרי ניקוד: {run.text}")
            
            # שמירת הקובץ
            logger.info(f"שומר קובץ פלט: {output_path}")
            target_doc.save(output_path)
            logger.info("הקובץ נשמר בהצלחה")
            
        except Exception as e:
            logger.error(f"שגיאה בעיבוד הקובץ: {str(e)}", exc_info=True)
            raise

    def test_known_dataset(self) -> None:
        """בדיקת המנקד על דאטה סט ידוע"""
        # דוגמה מוכרת עם ניקוד - עם שוליים נוספים
        source_text = """
        וַיְהִי בִּימֵי אֲחַשְׁוֵרוֹשׁ הוּא אֲחַשְׁוֵרוֹשׁ הַמֹּלֵךְ מֵהֹדּוּ וְעַד כּוּשׁ. בַּיָּמִים הָהָם כְּשֶׁבֶת הַמֶּלֶךְ אֲחַשְׁוֵרוֹשׁ עַל כִּסֵּא מַלְכוּתוֹ.
        בְּרֵאשִׁית בָּרָא אֱלֹהִים אֵת הַשָּׁמַיִם וְאֵת הָאָרֶץ. וְהָאָרֶץ הָיְתָה תֹהוּ וָבֹהוּ וְחֹשֶׁךְ עַל פְּנֵי תְהוֹם וְרוּחַ אֱלֹהִים מְרַחֶפֶת עַל פְּנֵי הַמָּיִם.
        וַיֹּאמֶר אֱלֹהִים יְהִי אוֹר וַיְהִי אוֹר. וַיַּרְא אֱלֹהִים אֶת הָאוֹר כִּי טוֹב וַיַּבְדֵּל אֱלֹהִים בֵּין הָאוֹר וּבֵין הַחֹשֶׁךְ.
        הַתְּפִלָּה הִיא עֲבוֹדַת הַלֵּב, וְהִיא עִקַּר הָעֲבוֹדָה שֶׁבַּלֵּב. בִּזְמַן שֶׁבֵּית הַמִּקְדָּשׁ הָיָה קַיָּם, הָיְתָה הָעֲבוֹדָה בְּקָרְבָּנוֹת.
        וְעַכְשָׁיו שֶׁאֵין לָנוּ לֹא כֹהֵן וְלֹא מִזְבֵּחַ, אֵין לָנוּ אֶלָּא תְּפִלָּה. וּתְפִלָּה זוֹ צְרִיכָה לִהְיוֹת בְּכַוָּנָה וּבְרָצוֹן.
        וְכָךְ אָמְרוּ חֲכָמִים: תְּפִלָּה בְּלֹא כַוָּנָה כְּגוּף בְּלֹא נְשָׁמָה. וְעוֹד אָמְרוּ: הַמִּתְפַּלֵּל צָרִיךְ שֶׁיְּכַוֵּן אֶת לִבּוֹ לַשָּׁמַיִם.
        """
        
        # טקסט יעד עם חלקים חופפים וחלקים שונים
        target_text = """
        בראשית ברא אלוהים את השמים ואת הארץ. והארץ היתה תוהו ובוהו וחושך על פני תהום ורוח אלוהים מרחפת על פני המים.
        ויאמר אלוהים יהי אור ויהי אור. וירא אלוהים את האור כי טוב ויבדל אלוהים בין האור ובין החושך.
        התפילה היא עבודת הלב, והיא עיקר העבודה שבלב. בזמן שבית המקדש היה קיים, היתה העבודה בקרבנות.
        ועכשיו שאין לנו לא כהן ולא מזבח, אין לנו אלא תפילה. ותפילה זו צריכה להיות בכוונה וברצון.
        וכך אמרו חכמים: תפילה בלא כוונה כגוף בלא נשמה. ועוד אמרו: המתפלל צריך שיכוון את לבו לשמים.
        """
        
        logger.info("\n=== בדיקת דאטה סט ידוע ===")
        
        # יצירת עותק וירטואלי
        source_nikud, source_clean = self._create_virtual_copy(source_text)
        
        # מציאת חפיפה
        start_s, end_s, start_t, end_t = self._find_overlap(source_clean, target_text)
        
        # בדיקת תוצאות
        if start_s == end_s == start_t == end_t == 0:
            logger.error("❌ נכשל: לא נמצאה חפיפה בדאטה סט ידוע")
            return
        
        result = self.add_nikud_to_text(source_text, target_text)
        
        # בדיקה שהתוצאה מכילה ניקוד
        if not self._has_nikud(result):
            logger.error("❌ נכשל: התוצאה לא מכילה ניקוד")
            return
        
        logger.info("✓ הצליח: נמצאה חפיפה והתוצאה מכילה ניקוד")
        logger.info(f"תוצאה:\n{result}")