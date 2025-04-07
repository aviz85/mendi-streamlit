"""
Nikud Service - Handle Hebrew vocalization (nikud) in text

This module provides functionality for handling Hebrew vocalization (nikud) marks.
It includes functions for removing nikud, adding nikud, and converting between
different forms of Hebrew text with or without nikud.
"""

import logging
from pathlib import Path
from typing import Tuple, Dict, List
from docx import Document
from docx.oxml import OxmlElement
import re

from .document_processor import DocumentProcessor
from .gemini_service import GeminiService
from .usage_logger import streamlit_logger as st_log

class NikudService:
    """
    Service for handling Hebrew vocalization (nikud) in documents
    
    This class provides methods for processing DOCX files with Hebrew text,
    handling vocalization marks, and merging content between files while
    preserving formatting and nikud.
    """
    
    def __init__(self, gemini_api_key: str = None):
        """
        Initialize the nikud service
        
        Parameters:
            gemini_api_key: Optional API key for Gemini AI service
        """
        self.logger = logging.getLogger(__name__)
        self.doc_processor = DocumentProcessor()
        self.gemini = GeminiService(api_key=gemini_api_key)
        
    def _read_docx(self, file_path: str) -> Tuple[str, Document]:
        """
        Read DOCX file and extract text while preserving bold formatting
        
        Parameters:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple containing the extracted text with HTML bold tags and the Document object
        """
        doc = Document(file_path)
        text = ""
        bold_count = 0
        
        st_log.log(f"קורא קובץ: {file_path}", "📖")
        
        for para in doc.paragraphs:
            para_text = ""
            current_bold_text = ""
            last_was_bold = False
            
            for run in para.runs:
                # Check only for Complex Script bold (w:bCs) in XML
                rPr = run._r.get_or_add_rPr()
                is_bold = bool(rPr.xpath('.//w:bCs'))
                run_text = run.text
                
                # Skip wrapping if it's just whitespace
                if run_text.strip() == "":
                    if last_was_bold and current_bold_text:
                        # Don't end bold section on whitespace - might continue
                        current_bold_text += run_text
                    else:
                        para_text += run_text
                    continue
                
                if is_bold:
                    if not last_was_bold:  # Start of bold section
                        current_bold_text = run_text
                    else:  # Continue bold section
                        current_bold_text += run_text
                    last_was_bold = True
                else:
                    if last_was_bold:  # End of bold section
                        # Only wrap if there's actual content
                        if current_bold_text.strip():
                            para_text += f"<b>{current_bold_text}</b>"
                            bold_count += 1
                        else:
                            para_text += current_bold_text
                        current_bold_text = ""
                    para_text += run_text
                    last_was_bold = False
            
            # Handle any remaining bold text at end of paragraph
            if last_was_bold and current_bold_text:
                # Only wrap if there's actual content
                if current_bold_text.strip():
                    para_text += f"<b>{current_bold_text}</b>"
                    bold_count += 1
                else:
                    para_text += current_bold_text
            
            text += para_text + "\n"
        
        st_log.log(f"זוהו {bold_count} קטעים מודגשים", "🔍")
        
        # Clean up multiple adjacent bold tags
        text = re.sub(r'</b>\s*<b>', '', text)
        
        return text, doc
    
    def _read_docx_raw(self, file_path: str) -> Tuple[List[str], Document]:
        """
        Read DOCX file and return raw paragraphs without HTML conversion
        
        Parameters:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple containing a list of raw paragraph texts and the Document object
        """
        doc = Document(file_path)
        paragraphs_raw = []
        
        st_log.log(f"קורא קובץ כפסקאות גולמיות: {file_path}", "📖")
        
        # Extract raw paragraphs
        for para in doc.paragraphs:
            paragraphs_raw.append(para.text)
            
        return paragraphs_raw, doc
        
    def _write_docx(self, content: str, template_doc: Document, output_path: str):
        """
        Write content back to DOCX preserving all formatting
        
        Parameters:
            content: Text content with HTML tags
            template_doc: Template document to copy styles from
            output_path: Path to save the output document
        """
        doc = Document()
        
        # Copy styles from template
        doc.styles._element = template_doc.styles._element
        
        # Set RTL at document level using XML
        doc._body._element.set('bidi', '1')
        
        # Process content and write to doc
        # Split by double newlines to separate paragraphs properly
        paragraphs = content.split('\n')
        
        for para_text in paragraphs:
            # Add empty paragraph to preserve line breaks
            if not para_text.strip():
                empty_para = doc.add_paragraph()
                empty_para.paragraph_format.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
                empty_para.style = template_doc.paragraphs[0].style if template_doc.paragraphs else None
                # Set paragraph direction
                pPr = empty_para._p.get_or_add_pPr()
                pPr.set('bidi', '1')
                continue
                
            para = doc.add_paragraph()
            # Set RTL and right alignment
            para.paragraph_format.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
            para.style = template_doc.paragraphs[0].style if template_doc.paragraphs else None
            
            # Set paragraph direction
            pPr = para._p.get_or_add_pPr()
            pPr.set('bidi', '1')
            
            # Split by bold markers - using non-greedy match for nested tags
            parts = re.split(r'(<b>.*?</b>)', para_text)
            
            for part in parts:
                if part.startswith('<b>') and part.endswith('</b>'):
                    # Bold text - extract content between tags
                    text = re.search(r'<b>(.*?)</b>', part).group(1)
                    run = para.add_run(text)
                    
                    # Set run direction and Complex Script bold only
                    rPr = run._r.get_or_add_rPr()
                    rPr.set('rtl', '1')
                    # Add w:bCs
                    bCs = OxmlElement('w:bCs')
                    rPr.append(bCs)
                else:
                    # Regular text
                    if part.strip() or part == " ":  # Preserve all text, including whitespace-only
                        run = para.add_run(part)
                        # Set run direction
                        rPr = run._r.get_or_add_rPr()
                        rPr.set('rtl', '1')
        
        # Save with error handling
        try:
            doc.save(output_path)
            st_log.log(f"המסמך נשמר בהצלחה: {output_path}", "💾")
        except Exception as e:
            st_log.log(f"שגיאה בשמירת המסמך: {str(e)}", "❌")
            raise

    def _convert_paragraphs_to_html_formatted_text(self, paragraphs: List[str], doc: Document) -> str:
        """
        Convert raw paragraphs to HTML formatted text with bold tags
        
        Parameters:
            paragraphs: List of raw paragraph texts
            doc: Document object containing formatting information
            
        Returns:
            HTML formatted text with bold tags
        """
        html_formatted_text = ""
        bold_count = 0
        
        # Map paragraph indices to the actual paragraph objects
        doc_paragraphs = doc.paragraphs
        
        for i, para_text in enumerate(paragraphs):
            if i < len(doc_paragraphs):  # Safety check
                para = doc_paragraphs[i]
                html_para = ""
                
                # Skip if paragraph has no content
                if not para_text.strip():
                    html_formatted_text += "\n"
                    continue
                
                # Track start and end positions for each run
                position = 0
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue
                    
                    # Check if bold
                    rPr = run._r.get_or_add_rPr()
                    is_bold = bool(rPr.xpath('.//w:bCs'))
                    
                    if is_bold and run_text.strip():
                        html_para += f"<b>{run_text}</b>"
                        bold_count += 1
                    else:
                        html_para += run_text
                    
                    position += len(run_text)
                
                html_formatted_text += html_para + "\n"
            else:
                # Fallback if we somehow have more paragraphs than in the document
                html_formatted_text += para_text + "\n"
        
        st_log.log(f"זוהו {bold_count} קטעים מודגשים", "🔍")
        
        # Clean up multiple adjacent bold tags
        html_formatted_text = re.sub(r'</b>\s*<b>', '', html_formatted_text)
        
        return html_formatted_text

    def process_files(self, source_path: str, target_path: str, output_path: str, create_debug_file: bool = False):
        """
        Process source and target files to add nikud
        
        This method processes a source document (without nikud) and a target document 
        (with nikud), identifies matching sections, and creates a new document with 
        the formatting of the source document but with nikud from the target document.
        
        Parameters:
            source_path: Path to the source file (without nikud)
            target_path: Path to the target file (with nikud)
            output_path: Path to save the processed file
            create_debug_file: Whether to create a debug file with additional information
        """
        st_log.log("מתחיל תהליך הוספת ניקוד", "🚀")
        
        # Create report file
        report_path = output_path.replace('.docx', '_report.txt')
        with open(report_path, 'w', encoding='utf-8') as report_file:
            report_file.write(f"דוח עיבוד ניקוד\n")
            report_file.write(f"==============\n\n")
            report_file.write(f"קובץ מקור: {source_path}\n")
            report_file.write(f"קובץ יעד: {target_path}\n")
            report_file.write(f"קובץ פלט: {output_path}\n\n")
        
        # Create debug file if requested
        debug_path = None
        if create_debug_file:
            debug_path = output_path.replace('.docx', '_debug.txt')
            with open(debug_path, 'w', encoding='utf-8') as debug_file:
                debug_file.write(f"קובץ דיבוג לתהליך הניקוד\n")
                debug_file.write(f"======================\n\n")
        
        # Read files - first as raw paragraphs
        st_log.log("קורא קבצים...", "📂")
        source_paragraphs, source_doc = self._read_docx_raw(source_path)
        target_paragraphs, target_doc = self._read_docx_raw(target_path)
        
        # Log paragraph counts for debugging
        self._append_to_report(report_path, f"נמצאו {len(source_paragraphs)} פסקאות בקובץ המקור\n")
        self._append_to_report(report_path, f"נמצאו {len(target_paragraphs)} פסקאות בקובץ היעד\n\n")
        
        # First process source documents - split to sections based on raw paragraphs
        st_log.log("מפצל את המסמך המקורי לחלקים...", "📑")
        source_sections = self.doc_processor.split_to_sections_from_raw(source_paragraphs)
        
        st_log.log("מפצל את מסמך היעד לחלקים...", "📑")
        target_sections = self.doc_processor.split_to_sections_from_raw(target_paragraphs)
        
        # Now convert to HTML format for further processing
        source_text = self._convert_paragraphs_to_html_formatted_text(source_paragraphs, source_doc)
        target_text = self._convert_paragraphs_to_html_formatted_text(target_paragraphs, target_doc)
        
        # Add to report
        self._append_to_report(report_path, f"נמצאו {len(source_sections)} חלקים במקור\n")
        self._append_to_report(report_path, f"נמצאו {len(target_sections)} חלקים ביעד\n\n")
        
        # Find matching sections
        matches = self.doc_processor.find_matching_sections(source_sections, target_sections)
        self._append_to_report(report_path, f"נמצאו {len(matches)} התאמות בין חלקי המסמכים\n\n")
        
        # Process each match with Gemini
        st_log.log("מעבד חלקים...", "⚙️")
        processed_sections = {}
        
        # Create a new GeminiService instance for each section to avoid context accumulation
        for idx, (source_section, target_section) in enumerate(matches):
            st_log.log(f"מעבד חלק {idx+1} מתוך {len(matches)}: {target_section.header}", "⚙️")
            
            content = self.doc_processor.prepare_for_nikud(source_section, target_section)
            if not content:
                continue
            
            self._append_to_report(report_path, f"מעבד חלק: {target_section.header}\n")
            self._append_to_report(report_path, f"נמצאו {content.get('bold_words_count', 0)} מילים/קטעים מודגשים בחלק זה\n")
            
            # Add debug information if requested
            if debug_path:
                self._append_to_debug(debug_path, f"\n\n{'='*50}\n")
                self._append_to_debug(debug_path, f"חלק: {target_section.header}\n")
                self._append_to_debug(debug_path, f"{'='*50}\n\n")
                self._append_to_debug(debug_path, "טקסט מקור (מנוקד):\n\n")
                self._append_to_debug(debug_path, f"{content['source_content']}\n\n")
                self._append_to_debug(debug_path, "טקסט יעד (עם תגי <b>):\n\n")
                self._append_to_debug(debug_path, f"{content['target_content']}\n\n")
            
            # Create a fresh Gemini instance for each section to avoid context issues
            section_gemini = GeminiService(api_key=self.gemini.api_key)
            processed_content = section_gemini.add_nikud(content, report_path)
            
            # Add debug information for processed content
            if debug_path:
                self._append_to_debug(debug_path, "תוצאה (אחרי עיבוד):\n\n")
                self._append_to_debug(debug_path, f"{processed_content}\n\n")
                
            # Save the processed content
            processed_sections[target_section.header] = processed_content
            
        # Recombine the processed sections
        st_log.log("מרכיב מסמך מעובד...", "📄")
        final_content = []
        for section in target_sections:
            if section.header in processed_sections:
                final_content.append(processed_sections[section.header])
            else:
                # If section wasn't processed, use original content
                final_content.append(section.content)
        
        # Write the combined content to the output file
        self._write_docx('\n'.join(final_content), target_doc, output_path)
        
        # Log completion
        st_log.log(f"תהליך הוספת הניקוד הושלם. קובץ נשמר ב: {output_path}", "✅")

    def _append_to_report(self, report_path: str, text: str):
        """
        Append text to the report file
        
        Parameters:
            report_path: Path to the report file
            text: Text to append
        """
        with open(report_path, 'a', encoding='utf-8') as report_file:
            report_file.write(text)

    def _append_to_debug(self, debug_path: str, text: str):
        """
        Append text to the debug file
        
        Parameters:
            debug_path: Path to the debug file
            text: Text to append
        """
        with open(debug_path, 'a', encoding='utf-8') as debug_file:
            debug_file.write(text)

    def _validate_nikud_in_bold_only(self, text: str) -> bool:
        """
        Validate that nikud appears only in bold text
        
        Parameters:
            text: Text to validate
            
        Returns:
            True if nikud appears only in bold text, False otherwise
        """
        # Define nikud pattern
        nikud_pattern = r'[\u05B0-\u05BC\u05C1-\u05C2\u05C4-\u05C5\u05C7]'
        
        # Split text by bold tags
        parts = re.split(r'(<b>.*?</b>)', text)
        
        for part in parts:
            if not part.startswith('<b>') and re.search(nikud_pattern, part):
                # Found nikud in non-bold text
                return False
            
        return True

    def _fix_nikud_overflow(self, text: str) -> str:
        """
        Fix text by removing nikud from non-bold parts
        
        Parameters:
            text: Text to fix
            
        Returns:
            Fixed text with nikud only in bold parts
        """
        # Define nikud pattern
        nikud_pattern = r'[\u05B0-\u05BC\u05C1-\u05C2\u05C4-\u05C5\u05C7]'
        
        # Split text by bold tags
        parts = re.split(r'(<b>.*?</b>)', text)
        
        result = []
        for part in parts:
            if part.startswith('<b>'):
                # Keep bold parts as is
                result.append(part)
            else:
                # Remove nikud from non-bold parts
                clean_part = re.sub(nikud_pattern, '', part)
                result.append(clean_part)
            
        # Join back without removing any whitespace or newlines
        result_text = ''.join(result)
        
        return result_text

    def add_nikud(self, text: str) -> str:
        """
        Add nikud to Hebrew text
        
        Parameters:
            text: Hebrew text without nikud
            
        Returns:
            Hebrew text with nikud added
        """
        words = text.split()
        result = []
        
        for word in words:
            result.append(self.dummy_mappings.get(word, word))
            
        return " ".join(result)

    def remove_nikud(self, text: str) -> str:
        """
        Remove nikud from Hebrew text
        
        Parameters:
            text: Hebrew text with nikud
            
        Returns:
            Hebrew text without nikud
        """
        # Unicode ranges for nikud marks
        nikud_pattern = r'[\u05B0-\u05BC\u05C1-\u05C2\u05C4-\u05C5\u05C7]'
        import re
        return re.sub(nikud_pattern, '', text)

    def test(self) -> Dict[str, bool]:
        """
        Run basic tests with dummy data
        
        Returns:
            Dictionary with test results
        """
        test_cases = {
            "Basic word": {
                "input": "שלום",
                "expected": "שָׁלוֹם"
            },
            "Multiple words": {
                "input": "ברוך הבא",
                "expected": "בָּרוּךְ הַבָּא"
            }
        }
        
        results = {}
        for name, case in test_cases.items():
            result = self.add_nikud(case["input"])
            results[name] = result == case["expected"]
            
        return results 