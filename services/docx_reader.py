"""
DocxReader - Class for reading and analyzing DOCX files

This class provides functionality for reading a DOCX file and converting it to various formats,
including a raw paragraph list or styled HTML that preserves formatting.
"""
import logging
import re
from typing import List, Tuple, Dict, Any, Union

from docx import Document

from .usage_logger import streamlit_logger as st_log

class DocxReader:
    """Class for reading DOCX documents and converting them to different formats"""
    
    def __init__(self):
        """Initialize a new DOCX reader"""
        self.logger = logging.getLogger(__name__)
        
    def read_docx_raw(self, file_path: str) -> Tuple[List[str], Document]:
        """
        Read a DOCX file and return a list of raw paragraphs and the original Document object
        
        Parameters:
            file_path: Path to the DOCX file
            
        Returns:
            List of paragraphs and the original Document
        """
        st_log.log(f"קורא קובץ DOCX: {file_path}", "📖")
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)
                    
            st_log.log(f"נקראו {len(paragraphs)} פסקאות מהקובץ", "✓")
            return paragraphs, doc
            
        except Exception as e:
            st_log.log(f"שגיאה בקריאת הקובץ: {str(e)}", "❌")
            raise
            
    def read_docx_html(self, file_path: str) -> Tuple[str, int]:
        """
        Read a DOCX file and convert it to HTML while identifying and counting bold segments
        
        Parameters:
            file_path: Path to the DOCX file
            
        Returns:
            HTML content and bold segment count
        """
        st_log.log(f"קורא ומנתח פורמט DOCX: {file_path}", "🔍")
        
        try:
            # Read document
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para)
                
            # Convert to HTML and count bold segments
            html_content, bold_count = self.convert_paragraphs_to_html(paragraphs, doc)
            
            st_log.log(f"נמצאו {bold_count} קטעים מודגשים בקובץ", "🔤")
            return html_content, bold_count
            
        except Exception as e:
            st_log.log(f"שגיאה בקריאת הקובץ והמרתו ל-HTML: {str(e)}", "❌")
            raise
            
    def convert_paragraphs_to_html(self, paragraphs: List[Any], doc: Document) -> Tuple[str, int]:
        """
        Convert a list of paragraphs to HTML format with <b> tags for emphasis and count bold segments
        
        Parameters:
            paragraphs: List of paragraphs from the Document
            doc: The original Document object
            
        Returns:
            HTML string and the total number of bold segments
        """
        html_parts = []
        bold_count = 0
        
        for para in paragraphs:
            if not para.text.strip():
                continue
                
            # Process paragraph runs to detect formatting
            para_parts = []
            for run in para.runs:
                text = run.text
                if text.strip():  # Skip empty runs
                    if run.bold:
                        para_parts.append(f"<b>{text}</b>")
                        bold_count += 1
                    else:
                        para_parts.append(text)
                        
            # Join paragraph parts and add to html parts
            html_parts.append("".join(para_parts))
            
        # Join paragraphs with newlines
        html_content = "\n".join(html_parts)
        
        return html_content, bold_count 