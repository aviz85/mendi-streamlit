"""
DocxWriter - Class for writing DOCX files while preserving formatting

This class provides functionality for writing textual content back to DOCX files
while preserving formatting and supporting HTML tags like <b> for bold text
"""
import logging
import re
from typing import List
from docx import Document
from docx.oxml import OxmlElement
from .usage_logger import streamlit_logger as st_log

class DocxWriter:
    def __init__(self):
        """Create a new DOCX writer"""
        self.logger = logging.getLogger(__name__)
    
    def write_docx(self, content: str, template_doc: Document, output_path: str):
        """
        Write textual content to a DOCX file while preserving formatting
        
        Parameters:
            content: Text to write (can include HTML tags like <b>)
            template_doc: DOCX document to source styles from
            output_path: Path to save the new file
        """
        doc = Document()
        
        # Copy styles from template
        doc.styles._element = template_doc.styles._element
        
        # Set RTL at document level using XML
        doc._body._element.set('bidi', '1')
        
        # Process content and write to doc
        # Split by newlines to separate paragraphs properly
        paragraphs = content.split('\n')
        
        for para_text in paragraphs:
            # Add empty paragraph to preserve line breaks
            if not para_text.strip():
                empty_para = doc.add_paragraph()
                empty_para.paragraph_format.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
                empty_para.style = template_doc.paragraphs[0].style if template_doc.paragraphs else None
                # Set paragraph direction
                pPr = empty_para._p.get_or_add_pPr()
                pPr.set('bidi', '1')
                continue
                
            para = doc.add_paragraph()
            # Set RTL and right alignment
            para.paragraph_format.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
            para.style = template_doc.paragraphs[0].style if template_doc.paragraphs else None
            
            # Set paragraph direction
            pPr = para._p.get_or_add_pPr()
            pPr.set('bidi', '1')
            
            # Split by bold markers - using non-greedy match for nested tags
            parts = re.split(r'(<b>.*?</b>)', para_text)
            
            for part in parts:
                if part.startswith('<b>') and part.endswith('</b>'):
                    # Bold text - extract content between tags
                    text = re.search(r'<b>(.*?)</b>', part).group(1)
                    run = para.add_run(text)
                    
                    # Set run direction and Complex Script bold only
                    rPr = run._r.get_or_add_rPr()
                    rPr.set('rtl', '1')
                    # Add w:bCs
                    bCs = OxmlElement('w:bCs')
                    rPr.append(bCs)
                else:
                    # Regular text
                    if part.strip() or part == " ":  # Preserve all text, including whitespace-only
                        run = para.add_run(part)
                        # Set run direction
                        rPr = run._r.get_or_add_rPr()
                        rPr.set('rtl', '1')
        
        # Save with error handling
        try:
            doc.save(output_path)
            st_log.log(f"המסמך נשמר בהצלחה: {output_path}", "💾")
        except Exception as e:
            st_log.log(f"שגיאה בשמירת המסמך: {str(e)}", "❌")
            raise 