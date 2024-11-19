import pytest
from services.docx_generator import create_interpretation_docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

@pytest.fixture
def sample_interpretation():
    return {
        "letter": "א",
        "original_text": "אברהם אבינו הלך לעקידה",
        "difficult_words": [
            {"word": "עקידה", "explanation": "קשירה והקרבה"},
            {"word": "אבינו", "explanation": "האבא של עם ישראל"}
        ],
        "detailed_interpretation": [
            {
                "quote": "אברהם אבינו",
                "explanation": "מייסד האומה היהודית"
            },
            {
                "quote": "הלך לעקידה",
                "explanation": "הלך להקריב את בנו יצחק"
            }
        ]
    }

def test_create_interpretation_docx(sample_interpretation):
    doc = create_interpretation_docx(sample_interpretation)
    paragraphs = doc.paragraphs
    
    # Test letter paragraph
    assert paragraphs[0].text == "א"
    assert paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Test original text
    assert paragraphs[1].text == "אברהם אבינו הלך לעקידה"
    assert paragraphs[1].alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT

def test_rtl_and_bold_formatting(sample_interpretation):
    doc = create_interpretation_docx(sample_interpretation)
    
    # Test RTL for all paragraphs
    for para in doc.paragraphs:
        if para.text:  # Skip empty paragraphs
            pPr = para._p.get_or_add_pPr()
            assert pPr.xpath('.//w:bidi')
    
    # Test bold formatting in difficult words
    words_para = doc.paragraphs[4]
    bold_words = ["עקידה", "אבינו"]
    for run in words_para.runs:
        if any(word in run.text for word in bold_words):
            assert run.bold

def test_document_defaults(sample_interpretation):
    doc = create_interpretation_docx(sample_interpretation)
    
    # Test document direction
    section = doc.sections[0]
    bidi_elements = section._sectPr.xpath('./w:bidi')
    assert bidi_elements
    assert bidi_elements[0].get(qn('w:val')) == '1'
    
    # Test default font
    style = doc.styles['Normal']
    assert style.font.name == 'David'
    assert style.font.size.pt == 12