from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
import re
import os
import pytest
from services.nikud_mapper import NikudMapper
from services.nikud_service import NikudService
from services.document_processor import DocumentProcessor

@pytest.fixture
def mapper():
    return NikudMapper(bold_only=True)

class TestNikudMapper:
    @pytest.fixture
    def mapper(self):
        return NikudMapper()

    def test_create_virtual_copy(self, mapper):
        """בדיקת יצירת עותק וירטואלי של מילים מנוקדות"""
        text = "בְּרֵאשִׁית בָּרָא אֱלֹהִים"
        nikud_words, clean_words = mapper._create_virtual_copy(text)
        
        assert len(nikud_words) == 3
        assert len(clean_words) == 3
        assert nikud_words[0] == "בְּרֵאשִׁית"
        assert clean_words[0] == "בראשית"

    def test_real_content(self, mapper):
        """בדיקת תוכן אמיתי מהמסמכים"""
        source = "הַקֹּדֶשׁ הָעַצְמִי הוּא הַחָכְמָה, הַהַכָּרָה לְבַדָּהּ, זֹהַר הָאִידֵאוֹת הָעֶלְיוֹנוֹת, בִּהְיוֹתָן בְּרוּם עֶלְיוֹנִיּוּתָן, בְּלֹא שׁוּם הַגְבָּלָה, בְּלֹא שׁוּם פְּרָטִיּוּת שֶׁל צִבְיוֹן, אֲפִלּוּ שֶׁל כְּלָלֵי כְּלָלִיּוּת אֲשֶׁר בַּמְּצִיאוּת. שֶׁהַכֹּל הוּא פְּרָטִי וּמֻגְבָּל לְגַבֵּי הָרֹחַב וְאִי הַגְּבוּלִיּוּת שֶׁל הַחָכְמָה הָעֶלְיוֹנָה, בְּרוּם עִזּוּזָה."
        source = ' '.join(source.replace('\n', ' ').split())
        
        target = "הקודש העצמי הוא החכמה, ההכרה לבדה, זוהר האידיאות העליונות, בהיותן ברום עליוניותן, בלא שום הגבלה, בלא שום פרטיות של צביון, אפילו של כללי כלליות אשר במציאות. שהכל הוא פרטי ומוגבל לגבי הרוחב ואי הגבוליות של החכמה העליונה, ברום עזוזה."
        target = ' '.join(target.replace('\n', ' ').split())
        
        # מציאת חפיפה
        source_words = re.findall(f"{mapper.nikud_patterns['hebrew']}+", source)
        print("\nSource words:", source_words)
        print("\nTarget text:", target)
        start_s, end_s, start_t, end_t = mapper._find_overlap(source_words, target)
        print("\nOverlap:", start_s, end_s, start_t, end_t)
        
        # בדיקה שנמצאה חפיפה
        assert start_s < end_s
        assert end_s - start_s >= 20  # לפחות 20 מילים
        
        # הוספת ניקוד
        result = mapper.add_nikud_to_text(source, target)
        
        # בדיקה שהתוצאה מכילה ניקוד
        assert mapper._has_nikud(result)
        
        # בדיקה שהמילים המנוקדות זהות למקור
        source_matches = list(re.finditer(f"({mapper.nikud_patterns['hebrew']}+)([^\u0590-\u05FF]*)", source))
        result_matches = list(re.finditer(f"({mapper.nikud_patterns['hebrew']}+)([^\u0590-\u05FF]*)", result))
        
        # מציאת המילים בחפיפה
        source_words = [m.group(1) for m in source_matches]
        result_words = [m.group(1) for m in result_matches]
        
        # בדיקה שהמילים בחפיפה זהות
        nikud_count = 0
        for i in range(len(result_words)):
            if mapper._has_nikud(result_words[i]):
                nikud_count += 1
                # מציאת המילה המתאימה במקור
                target_clean = mapper._strip_nikud(result_words[i])
                found = False
                for source_word in source_words:
                    if mapper._strip_nikud(source_word) == target_clean:
                        assert result_words[i] == source_word
                        found = True
                        break
                assert found, f"לא נמצאה מילה מנוקדת מתאימה במקור עבור {result_words[i]}"
        
        # בדיקה שיש מספיק מילים מנוקדות
        assert nikud_count >= 20, f"נמצאו רק {nikud_count} מילים מנוקדות"

    def test_find_overlap(self, mapper):
        """בדיקת מציאת חפיפה בין טקסטים"""
        source_words = ["בראשית", "ברא", "אלהים", "את", "השמים", "ואת", "הארץ"]
        target_text = "פירוש על בראשית ברא אלהים בתורה"
        
        start_s, end_s, start_t, end_t = mapper._find_overlap(source_words, target_text)
        
        assert start_s < end_s  # נמצאה חפיפה
        assert end_s - start_s >= 3  # החפיפה מכילה לפחות 3 מילים
        assert "בראשית" in ' '.join(source_words[start_s:end_s])

    def test_add_nikud_to_text(self, mapper):
        """בדיקת הוספת ניקוד לטקסט"""
        source = "בְּרֵאשִׁית בָּרָא אֱלֹהִים"
        target = "פירוש על בראשית ברא אלהים בתורה"
        
        result = mapper.add_nikud_to_text(source, target)
        
        # בודק שהמילים המשותפות קיבלו ניקוד
        assert "בְּרֵאשִׁית" in result
        assert "בָּרָא" in result
        assert "אֱלֹהִים" in result
        
        # בודק שהמילים שלא היו במקור נשארו ללא ניקוד
        assert "פירוש" in result
        assert "על" in result
        assert "בתורה" in result

    def test_process_docx(self, tmp_path):
        """בדיקת עיבוד קובץ Word שלם"""
        mapper = NikudMapper(bold_only=True)
        
        # יצירת קובץ מקור
        source_doc = Document()
        source_doc.add_paragraph("בְּרֵאשִׁית בָּרָא אֱלֹהִים")
        source_path = tmp_path / "source.docx"
        source_doc.save(source_path)
        
        # יצירת קובץ יעד
        target_doc = Document()
        p = target_doc.add_paragraph()
        p.add_run("פירוש על ")
        p.add_run("בראשית").bold = True
        p.add_run(" ועל ")
        p.add_run("ברא").bold = True
        target_path = tmp_path / "target.docx"
        target_doc.save(target_path)
        
        # יצירת קובץ פלט
        output_path = tmp_path / "output.docx"
        
        # עיבוד הקבצים
        mapper.process_docx(str(source_path), str(target_path), str(output_path))
        
        # בדיקת התוצאה
        output_doc = Document(output_path)
        output_text = ""
        for para in output_doc.paragraphs:
            for run in para.runs:
                if run.bold:
                    assert mapper._has_nikud(run.text)  # בודק שטקסט מודגש קיבל ניקוד
                else:
                    assert not mapper._has_nikud(run.text)  # בודק שטקסט רגיל לא קיבל ניקוד

    def test_edge_cases(self, mapper):
        """בדיקת מקרי קצה"""
        # טקסט ריק
        assert mapper.add_nikud_to_text("", "") == ""
        
        # טקסט ללא מילים משותפות
        source = "בְּרֵאשִׁית בָּרָא"
        target = "שלום עולם"
        assert mapper.add_nikud_to_text(source, target) == target
        
        # טקסט זהה לחלוטין
        source = "בְּרֵאשִׁית"
        target = "בראשית"
        result = mapper.add_nikud_to_text(source, target)
        assert mapper._has_nikud(result)
        assert mapper._strip_nikud(result) == target

    def test_multiple_matches(self, mapper):
        """בדיקת טיפול במספר התאמות"""
        source = "בְּרֵאשִׁית בָּרָא אֱלֹהִים בְּרֵאשִׁית בָּרָא"
        target = "בראשית ברא אלהים ... בראשית ברא"
        
        result = mapper.add_nikud_to_text(source, target)
        
        # בודק שכל ההופעות של המילים קיבלו ניקוד
        assert result.count("בְּרֵאשִׁית") == 2
        assert result.count("בָּרָא") == 2

    def test_real_document_structure(self, tmp_path):
        """Test processing a document with real-world structure:
        - Sections with Hebrew letters
        - Bold source text with nikud
        - Non-bold explanatory text
        - Multiple paragraphs per section
        """
        # Create source document with nikud
        source_doc = Document()
        
        # א. First section
        source_doc.add_paragraph("א")
        source_doc.add_paragraph("בְּרֵאשִׁית בָּרָא אֱלֹהִים אֵת הַשָּׁמַיִם וְאֵת הָאָרֶץ")
        source_doc.add_paragraph("וְהָאָרֶץ הָיְתָה תֹהוּ וָבֹהוּ וְחֹשֶׁךְ עַל פְּנֵי תְהוֹם")
        
        # ב. Second section
        source_doc.add_paragraph("ב")
        source_doc.add_paragraph("וַיֹּאמֶר אֱלֹהִים יְהִי אוֹר וַיְהִי אוֹר")
        source_doc.add_paragraph("וַיַּרְא אֱלֹהִים אֶת הָאוֹר כִּי טוֹב")
        
        source_path = tmp_path / "source.docx"
        source_doc.save(str(source_path))
        
        # Create target document with mixed bold/non-bold text
        target_doc = Document()
        
        # א. First section
        target_doc.add_paragraph("א")
        # Introduction paragraph
        target_doc.add_paragraph("פרק א - בריאת העולם")
        
        # Key terms paragraph
        p = target_doc.add_paragraph()
        p.add_run("מילים מרכזיות: ")
        run = p.add_run("בראשית")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - ההתחלה, ")
        run = p.add_run("ברא אלהים")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - הבריאה האלוהית")
        
        # Main explanation with bold source text
        p = target_doc.add_paragraph()
        run = p.add_run("בראשית ברא אלהים את השמים ואת הארץ")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - כאן מתואר מעשה הבריאה הראשון. ")
        run = p.add_run("והארץ היתה תהו ובהו")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - תיאור מצב הארץ בתחילת הבריאה.")
        
        # ב. Second section
        target_doc.add_paragraph("ב")
        p = target_doc.add_paragraph()
        run = p.add_run("ויאמר אלהים יהי אור ויהי אור")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - בריאת האור. ")
        run = p.add_run("וירא אלהים את האור כי טוב")
        run._r.get_or_add_rPr().append(OxmlElement('w:bCs'))  # Make bold
        p.add_run(" - הערכת הבריאה.")
        
        target_path = tmp_path / "target.docx"
        target_doc.save(str(target_path))
        
        # Process the documents
        output_path = tmp_path / "output.docx"
        mapper = NikudMapper(bold_only=True)
        mapper.process_docx(str(source_path), str(target_path), str(output_path))
        
        # Verify the output
        output_doc = Document(str(output_path))
        
        # Count nikud and bold words
        nikud_count = 0
        non_nikud_count = 0
        
        for para in output_doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue
                    
                is_bold = bool(run._r.xpath('.//w:bCs'))
                has_nikud = bool(re.search(mapper.nikud_patterns['nikud'], text))
                
                if is_bold:
                    # Bold text should have nikud
                    assert has_nikud, f"Bold text missing nikud: {text}"
                    nikud_count += len(text.split())
                else:
                    # Non-bold text should not have nikud
                    assert not has_nikud, f"Non-bold text has nikud: {text}"
                    non_nikud_count += len(text.split())
        
        # Verify we have enough words of each type
        assert nikud_count >= 8, f"Not enough nikud words: {nikud_count}"
        assert non_nikud_count >= 5, f"Not enough non-nikud words: {non_nikud_count}"

    def test_real_files(self):
        """Test processing real source.docx and target.docx files"""
        from services.nikud_service import NikudService
        from services.document_processor import DocumentProcessor
        import os
        
        # Initialize services with API key
        GEMINI_API_KEY = "AIzaSyCQCtxTshjHi3gWLI7MsIBwpkbokrZ2sx4"
        nikud_service = NikudService(gemini_api_key=GEMINI_API_KEY)
        doc_processor = DocumentProcessor()
        
        # Process the files
        source_path = "source.docx"
        target_path = "target.docx"
        output_path = "output.docx"
        
        try:
            # Read and process files
            print("\n=== קריאת קבצים ===")
            source_text, source_doc = nikud_service._read_docx(source_path)
            target_text, target_doc = nikud_service._read_docx(target_path)
            
            print("\n=== פיצול לחלקים ===")
            source_sections = doc_processor.split_to_sections(source_text)
            target_sections = doc_processor.split_to_sections(target_text)
            
            print(f"\nנמצאו {len(source_sections)} חלקים במקור")
            for section in source_sections:
                print(f"\nכותרת: {section.header}")
                print(f"תוכן ({len(section.content)} תווים):")
                print(section.content[:200] + "..." if len(section.content) > 200 else section.content)
            
            print(f"\nנמצאו {len(target_sections)} חלקים ביעד")
            for section in target_sections:
                print(f"\nכותרת: {section.header}")
                print(f"תוכן ({len(section.content)} תווים):")
                print(section.content[:200] + "..." if len(section.content) > 200 else section.content)
            
            print("\n=== מציאת התאמות ===")
            matches = doc_processor.find_matching_sections(source_sections, target_sections)
            print(f"נמצאו {len(matches)} התאמות")
            
            print("\n=== עיבוד חלקים ===")
            processed_sections = {}
            for source_section, target_section in matches:
                print(f"\nמעבד חלק: {target_section.header}")
                content = doc_processor.prepare_for_nikud(source_section, target_section)
                if content:
                    # Use NikudMapper for processing
                    mapper = NikudMapper(bold_only=True)
                    processed_content = mapper.add_nikud_to_text(content["source_content"], content["target_content"])
                    processed_sections[target_section.header] = processed_content
                    print(f"תוצאה ({len(processed_content)} תווים):")
                    print(processed_content[:200] + "..." if len(processed_content) > 200 else processed_content)
            
            print("\n=== הרכבת המסמך ===")
            final_content = []
            for section in target_sections:
                if section.header in processed_sections:
                    final_content.append(processed_sections[section.header])
                else:
                    final_content.append(section.content)
            
            print("\n=== שמירת המסמך ===")
            nikud_service._write_docx('\n'.join(final_content), target_doc, output_path)
            print(f"המסמך נשמר בהצלחה: {output_path}")
            
            # Verify output
            print("\n=== בדיקת הפלט ===")
            output_doc = Document(output_path)
            
            # Count nikud and bold words
            nikud_count = 0
            non_nikud_count = 0
            
            for para in output_doc.paragraphs:
                for run in para.runs:
                    text = run.text.strip()
                    if not text:
                        continue
                        
                    rPr = run._r.get_or_add_rPr()
                    is_bold = bool(rPr.xpath('.//w:bCs'))
                    has_nikud = bool(re.search(nikud_service.doc_processor.nikud_patterns['nikud'], text))
                    
                    if is_bold:
                        # Bold text should have nikud
                        assert has_nikud, f"טקסט מודגש ללא ניקוד: {text}"
                        nikud_count += len(text.split())
                    else:
                        # Non-bold text should not have nikud
                        assert not has_nikud, f"טקסט רגיל עם ניקוד: {text}"
                        non_nikud_count += len(text.split())
            
            print(f"\nסיכום:")
            print(f"מילים מנוקדות: {nikud_count}")
            print(f"מילים ללא ניקוד: {non_nikud_count}")
            
        finally:
            # Clean up
            if os.path.exists(output_path):
                os.remove(output_path)