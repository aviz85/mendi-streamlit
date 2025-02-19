import os
import logging
from pathlib import Path
import pytest
from docx import Document

from services.nikud_service import NikudService

# Configure logging
logging.basicConfig(level=logging.INFO)

@pytest.fixture
def service():
    return NikudService()

@pytest.fixture
def test_files(tmp_path):
    # Create test documents
    source_doc = Document()
    source_doc.add_paragraph("פרק א").bold = True
    source_doc.add_paragraph("בְּרֵאשִׁית בָּרָא אֱלֹהִים אֵת הַשָּׁמַיִם וְאֵת הָאָרֶץ")
    source_path = tmp_path / "source.docx"
    source_doc.save(source_path)
    
    target_doc = Document()
    target_doc.add_paragraph("פרק א").bold = True
    target_doc.add_paragraph("בראשית ברא אלהים את השמים ואת הארץ")
    target_path = tmp_path / "target.docx"
    target_doc.save(target_path)
    
    output_path = tmp_path / "output.docx"
    
    return source_path, target_path, output_path

def test_process_files(service, test_files):
    source_path, target_path, output_path = test_files
    
    # Process files
    service.process_files(str(source_path), str(target_path), str(output_path))
    
    # Verify output file exists
    assert output_path.exists()
    
    # Read output and verify content
    output_doc = Document(output_path)
    output_text = "\n".join(p.text for p in output_doc.paragraphs)
    
    # Basic verification
    assert "פרק א" in output_text
    assert "בְּרֵאשִׁית" in output_text  # Should contain nikud
    
    # Verify bold formatting preserved
    assert any(run.bold for para in output_doc.paragraphs for run in para.runs)

if __name__ == "__main__":
    pytest.main([__file__]) 