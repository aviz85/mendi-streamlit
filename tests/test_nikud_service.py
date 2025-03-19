import pytest
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from pathlib import Path
import os

class TestNikudSystem:
    @pytest.fixture
    def create_test_doc(self, tmp_path):
        def _create_doc(content_list, has_nikud=False, has_bold=False):
            doc = Document()
            for text in content_list:
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                if has_bold and has_nikud:
                    run = p.add_run(text)
                    run._element.get_or_add_rPr().append(parse_xml(r'<w:bCs/>'))
                else:
                    p.add_run(text)
            
            test_file = tmp_path / "test.docx"
            doc.save(test_file)
            return str(test_file)
        return _create_doc

    def test_paragraph_separation(self, create_test_doc):
        """בדיקת זיהוי פסקאות לוגיות"""
        content = [
            "פסקה ראשונה",
            "",  # שורה ריקה בודדת - לא אמורה להפריד
            "המשך פסקה ראשונה",
            "",
            "",  # שתי שורות ריקות - אמורות להפריד
            "פסקה שנייה"
        ]
        doc_path = create_test_doc(content)
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        # צריך להיות 2 פסקאות לוגיות
        assert len([p for p in result['paragraphs'] if p['text'].strip()]) == 2

    def test_nikud_detection(self, create_test_doc):
        """בדיקת זיהוי וספירת ניקוד"""
        content_with_nikud = ["הַשָּׁמַיִם", "הָאָרֶץ"]
        content_without_nikud = ["השמים", "הארץ"]
        
        doc_with_nikud = create_test_doc(content_with_nikud, has_nikud=True)
        doc_without_nikud = create_test_doc(content_without_nikud)
        
        analyzer = DocAnalyzer()
        result_with_nikud = analyzer.analyze_document(doc_with_nikud)
        result_without_nikud = analyzer.analyze_document(doc_without_nikud)
        
        assert result_with_nikud['total_nikud'] > 0
        assert result_without_nikud['total_nikud'] == 0

    def test_bold_formatting(self, create_test_doc):
        """בדיקת זיהוי טקסט מודגש"""
        content = ["טקסט רגיל", "טקסט מודגש"]
        doc_path = create_test_doc(content, has_bold=True)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        assert result['bold_stats']['complex_bold'] > 0

    def test_hebrew_word_count(self, create_test_doc):
        """בדיקת ספירת מילים בעברית"""
        content = [
            "חמש מילים בעברית בלבד",  # 5 מילים
            "Five words mixed with עברית ואנגלית",  # 2 מילים בעברית
        ]
        doc_path = create_test_doc(content)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        first_para = next(p for p in result['paragraphs'] if 'חמש' in p['text'])
        assert first_para['word_count'] == 5

    def test_significant_paragraphs(self, create_test_doc):
        """בדיקת זיהוי פסקאות משמעותיות"""
        short_text = "שלוש מילים קצרות"  # פחות מ-20 מילים
        long_text = " ".join(["מילה"] * 25)  # 25 מילים
        
        doc_path = create_test_doc([short_text, long_text])
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        assert result['significant_paragraphs'] == 1

    def test_mixed_formatting(self, create_test_doc):
        """בדיקת שילוב של ניקוד והדגשה"""
        content = ["הַטֶקְסְט הַמְנֻקָּד"]
        doc_path = create_test_doc(content, has_nikud=True, has_bold=True)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        assert result['paragraphs_with_nikud'] > 0
        assert result['bold_stats']['complex_bold'] > 0

    def test_empty_paragraphs(self, create_test_doc):
        """בדיקת טיפול בפסקאות ריקות"""
        content = ["", "טקסט רגיל", "", "", "טקסט נוסף", ""]
        doc_path = create_test_doc(content)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        # צריך להתעלם מפסקאות ריקות
        assert len([p for p in result['paragraphs'] if p['text'].strip()]) == 2

    def test_special_characters(self, create_test_doc):
        """בדיקת טיפול בתווים מיוחדים"""
        content = ["טקסט! עם? סימני& פיסוק@ מיוחדים#"]
        doc_path = create_test_doc(content)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        # צריך לספור מילים נכון גם עם סימני פיסוק
        first_para = result['paragraphs'][0]
        assert first_para['word_count'] == 5

    def test_rtl_handling(self, create_test_doc):
        """בדיקת טיפול בכיווניות טקסט"""
        content = ["טקסט מימין לשמאל", "Left to right text", "מעורב mixed טקסט"]
        doc_path = create_test_doc(content)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        # צריך לזהות רק מילים בעברית
        hebrew_words = sum(p['word_count'] for p in result['paragraphs'])
        assert hebrew_words == 4  # מימין + לשמאל + מעורב + טקסט

    def test_document_statistics(self, create_test_doc):
        """בדיקת סטטיסטיקות ברמת המסמך"""
        content = [
            "פסקה עם ניקוד הַטֶקְסְט",
            "פסקה בלי ניקוד",
            "פסקה עם הדגשה"
        ]
        doc_path = create_test_doc(content, has_nikud=True, has_bold=True)
        
        analyzer = DocAnalyzer()
        result = analyzer.analyze_document(doc_path)
        
        assert isinstance(result['total_paragraphs'], int)
        assert isinstance(result['significant_paragraphs'], int)
        assert isinstance(result['paragraphs_with_nikud'], int)
        assert isinstance(result['total_nikud'], int)
        assert all(isinstance(v, int) for v in result['bold_stats'].values()) 